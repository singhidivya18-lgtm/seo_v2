"""DOCX report generator tool.

Converts the article into a professionally formatted .docx file
that can be downloaded.
"""

import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


async def generate_docx(
    article_text: str,
    title: str = "Article",
    output_dir: str | None = None,
    job_id: str | None = None,
    linkedin_post: str | None = None,
    twitter_thread: str | None = None,
    image_paths: str | list[str] | None = None,
    approval_note: str | None = None,
) -> dict[str, object]:
    """Generate a downloadable .docx file from the article text.

    Use this tool when:
      - The user wants the article as a downloadable .docx document.
      - The user needs a formatted document for sharing or printing.

    Do NOT use this tool when:
      - The user only wants text output.
      - There is no article text available.

    Args:
        article_text: The full article text to convert.
        title: The document title. Defaults to "Article".
        output_dir: Directory to save the .docx file. Defaults to SEO_DATA_DIR
            (mounted volume in a container) or the pipeline root when unset.
        linkedin_post: Optional LinkedIn post text to include as a section.
        twitter_thread: Optional Twitter thread text to include as a section.
        image_paths: Optional image file path(s) to embed in the document.

    Returns:
        dict:
        {
            "status": "success",
            "platform": "docx",
            "filepath": str,
            "filename": str,
            "filesize_bytes": int
        }
        or {"status": "error", "error_message": str}
    """
    if not article_text or len(article_text.strip()) < 10:
        return {"status": "error", "error_message": "Article text is too short to generate a document."}

    if output_dir is None:
        output_dir = (
            os.environ.get("SEO_DATA_DIR", "").strip()
            or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        )

    try:
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        title_heading = doc.add_heading(title, level=0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if approval_note:
            note = doc.add_paragraph()
            note_run = note.add_run(approval_note)
            note_run.font.size = Pt(11)
            note_run.font.color.rgb = RGBColor(255, 0, 0)
            note_run.bold = True
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        timestamp_paragraph = doc.add_paragraph()
        timestamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = timestamp_paragraph.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_page_break()

        lines = article_text.split("\n")
        in_code_block = False
        code_lines = []

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code_block:
                    if code_lines:
                        code_text = "\n".join(code_lines)
                        p = doc.add_paragraph()
                        run = p.add_run(code_text)
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)
                    code_lines = []
                in_code_block = not in_code_block
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if not stripped:
                doc.add_paragraph()
                continue

            if stripped.startswith("# "):
                doc.add_heading(stripped[2:].strip(), level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip(), level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                text = stripped[2:].strip()
                p = doc.add_paragraph(text, style="List Bullet")
            elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in (".", ")"):
                text = stripped[2:].strip()
                p = doc.add_paragraph(text, style="List Number")
            elif stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(stripped[2:-2])
                run.bold = True
            elif stripped.startswith("---") or stripped.startswith("***"):
                doc.add_paragraph("_" * 40)
            else:
                p = doc.add_paragraph(stripped)

        if linkedin_post or twitter_thread or image_paths:
            doc.add_page_break()
            doc.add_heading("Social Media Versions", level=1)

        if linkedin_post:
            doc.add_heading("LinkedIn Post", level=2)
            for line in str(linkedin_post).split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

        if twitter_thread:
            doc.add_heading("Twitter Thread", level=2)
            for i, line in enumerate(str(twitter_thread).split("\n"), start=1):
                line = line.strip()
                if line:
                    doc.add_paragraph(f"{i}/ {line}")

        if image_paths:
            doc.add_heading("Images", level=2)
            if isinstance(image_paths, str):
                image_paths = [image_paths]
            for idx, img_path in enumerate(image_paths, start=1):
                img_path = str(img_path).strip()
                try:
                    if os.path.isfile(img_path):
                        doc.add_picture(img_path, width=Inches(5.5))
                        caption = doc.add_paragraph()
                        cap_run = caption.add_run(f"Image {idx}")
                        cap_run.font.size = Pt(9)
                        cap_run.font.color.rgb = RGBColor(128, 128, 128)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        note = doc.add_paragraph()
                        note_run = note.add_run(f"Image {idx} not found: {os.path.basename(img_path)}")
                        note_run.font.size = Pt(9)
                        note_run.font.color.rgb = RGBColor(255, 0, 0)
                except Exception as img_err:
                    note = doc.add_paragraph()
                    note_run = note.add_run(f"Image {idx} could not be embedded: {img_err}")
                    note_run.font.size = Pt(9)
                    note_run.font.color.rgb = RGBColor(255, 0, 0)

        os.makedirs(output_dir, exist_ok=True)
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        # URL-safe filename: keep [A-Za-z0-9_-] only; spaces/slashes -> "-", collapse
        safe_title = re.sub(r"[^A-Za-z0-9_-]+", "-", title).strip("-")[:50]
        if job_id:
            # structural link: batch caller asserts `job_id in filename`
            filename = f"{safe_title}_{job_id}_{timestamp_str}.docx"
        else:
            filename = f"{safe_title}_{timestamp_str}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        filesize = os.path.getsize(filepath)

        return {
            "status": "success",
            "platform": "docx",
            "filepath": filepath,
            "filename": filename,
            "filesize_bytes": filesize,
        }

    except Exception as e:
        return {"status": "error", "error_message": f"DOCX generation failed: {str(e)}"}